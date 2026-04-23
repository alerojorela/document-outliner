# -*- coding: utf-8 -*-
# by Alejandro Rojo Gualix 2022-02 ...
__author__ = 'Alejandro Rojo Gualix'
"""
pip install python-docx freeplane-io
"""

import subprocess
import sys
import tempfile
import argparse
from itertools import chain
from pathlib import Path

import re
import freeplane
# python-docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

"""
SR. No.	Colour Name In WD_COLOR_INDEX	Colour Description
1.	AUTO	Default or No Colour
2.	BLACK	Black Colour
3.	BLUE	Blue Colour
4.	BRIGHT_GREEN	Green Colour
5.	DARK_BLUE	Dark Blue Colour
6.	DARK_RED	Dark Red Colour
7.	DARK_YELLOW	Dark Yellow Colour
8.	GRAY_25	Light Gray Colour
9.	GRAY_50	Dark Gray Colour
10.	GREEN	Dark Green Colour
11.	PINK	Magenta Colour
12.	RED	Red Colour
13.	TEAL	Dark Cyan Colour
14.	TURQUOISE	Cyan Colour
15.	VIOLET	Dark Magenta Colour
16.	WHITE	White Colour
17.	YELLOW	Yellow Colour
"""

extensions = ['*.odt', '*.docx']
extensions = ['*.docx']

_BRIDGE_CHARS = set(' \t\n.,;:!?…-–—')


def _build_segments(runs, text_filter, paragraph):
    """Collect selected runs into segments.

    Non-selected runs that contain only whitespace / punctuation are treated as
    bridges: their text is kept and merged into the ongoing segment.  A
    non-selected run with substantial content closes the current segment.
    """
    segments = []
    buf = []
    bridge = ""

    for run in runs:
        if text_filter(run, paragraph):
            buf.append(bridge + run.text)
            bridge = ""
        else:
            if buf:
                if all(c in _BRIDGE_CHARS for c in run.text):
                    bridge += run.text   # hold it; include only if a selected run follows
                else:
                    segments.append("".join(buf))
                    buf = []
                    bridge = ""
            # if no segment is open, skip unselected runs entirely

    if buf:
        segments.append("".join(buf))

    return segments


class Outliner:
    def __init__(self, file: Path):
        self.file = file
        # look for title?
        self.name = file.stem
        self.document = Document(file)
        self._output_document = Document()
        self._output_midmapping = None

    def extract_marks(self, doc_file: Path = None, freeplane_file: Path = None, summarizer=None,
                      select_bold=True, select_italic=False, select_underline=False,
                      highlighted_color=None, markdown_mode=False, single_prompt=False):
        md_out = doc_file and Path(doc_file).suffix == '.md'
        md_lines = [] if md_out else None

        sample = ''
        for paragraph in self.document.paragraphs:
            sample += paragraph.text + '\n'
            if len(sample) > 100:
                break
        if not sample:
            return

        def effective(run, paragraph, attr):
            """Resolve formatting walking up the style hierarchy (run → char style → para style → base)."""
            val = getattr(run.font, attr)
            if val is not None:
                return val
            style = paragraph.style
            while style:
                v = getattr(style.font, attr, None)
                if v is not None:
                    return v
                style = style.base_style
            return False

        def text_filter(run, paragraph):
            return (select_bold      and effective(run, paragraph, 'bold'))      or \
                   (select_italic    and effective(run, paragraph, 'italic'))    or \
                   (select_underline and effective(run, paragraph, 'underline')) or \
                   (highlighted_color and run.font.highlight_color == highlighted_color)

        parents_stack = []
        if freeplane_file:
            self._output_midmapping = freeplane.Mindmap()
            self._output_midmapping.rootnode.plaintext = self.name
            parents_stack = [self._output_midmapping.rootnode]

        def append_summary(text, heading=""):
            summary = summarizer(text)  # summarizer is a callable: str -> str

            if doc_file:
                if md_out:
                    if heading:
                        md_lines.append(f"## {heading}\n")
                    md_lines.append(summary + "\n")
                else:
                    self._output_document.add_paragraph(summary)
            if freeplane_file:  # Freeplane
                node = parents_stack[-1].add_child(summary)
                node._node.attrib["STYLE"] = 'fork'

        def paragraph_to_md(p) -> str:
            result = ""
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                result += t
            return result

        if summarizer and single_prompt:
            lines = []
            for paragraph in self.document.paragraphs:
                style = paragraph.style.name
                if style.startswith("Title") or style.startswith("Heading"):
                    lines.append(f"# {paragraph.text}")
                else:
                    lines.append(paragraph_to_md(paragraph) if markdown_mode else paragraph.text)
            full_text = '\n'.join(lines)
            result = summarizer(full_text)
            if doc_file:
                if md_out:
                    Path(doc_file).write_text(result, encoding='utf-8')
                else:
                    self._output_document.add_paragraph(result)
                    self._output_document.save(doc_file)
            if freeplane_file:
                node = parents_stack[-1].add_child(result)
                node._node.attrib["STYLE"] = 'fork'
                self._output_midmapping.save(freeplane_file, encoding='utf-8')
            return

        styles = set()
        section_text = []
        current_heading = ""
        for paragraph in self.document.paragraphs:
            style = paragraph.style.name
            styles.add(style)

            if style.startswith("Title") or style.startswith("Heading"):
                if summarizer and section_text:
                    prefix = f"# {current_heading}\n\n" if markdown_mode and current_heading else ""
                    compiled_text = prefix + '\n'.join(section_text)
                    append_summary(compiled_text, heading=current_heading)
                    section_text = []
                current_heading = paragraph.text

                if style.startswith("Title"):
                    level = 1  # TODO 0
                elif style.startswith("Heading"):
                    m = re.search(r'(\d+)$', style)
                    level = int(m.group(0))

                if doc_file:
                    if md_out:
                        md_lines.append(f"{'#' * level} {paragraph.text}\n")
                    else:
                        self._output_document.add_heading(paragraph.text, level)

                if freeplane_file:
                    # Freeplane
                    # print('\t', level, len(parents_stack) - 1)
                    if level > len(parents_stack) - 1:  # rise level
                        while level > len(parents_stack):
                            parents_stack.append(parents_stack[-1].add_child('<missing branch>'))
                    else:  # drop: same or lower the level
                        while len(parents_stack) - level > 0:
                            parents_stack.pop()
                    node = parents_stack[-1].add_child(paragraph.text)
                    node._node.attrib["STYLE"] = 'bubble'
                    parents_stack.append(node)

            # elif paragraph.style.name == "Normal":
            # elif style.startswith("Body Text"):
            else:  # Body Text
                if summarizer:
                    # accumulate whole section
                    section_text.append(paragraph_to_md(paragraph) if markdown_mode else paragraph.text)
                else:
                    selected_text = _build_segments(paragraph.runs, text_filter, paragraph)
                    if selected_text:
                        if doc_file:
                            compiled_text = ' […] '.join(selected_text)
                            self._output_document.add_paragraph(compiled_text)
                        if freeplane_file:  # Freeplane
                            # TODO: use topic models
                            node_parent = parents_stack[-1].add_child(selected_text[0])
                            node_parent._node.attrib["STYLE"] = 'fork'
                            for segment in selected_text[1:]:
                                node = node_parent.add_child(segment)
                                node._node.attrib["STYLE"] = 'fork'

        # summarize residual last section if needed
        if summarizer and section_text:
            prefix = f"# {current_heading}\n\n" if markdown_mode and current_heading else ""
            compiled_text = prefix + '\n'.join(section_text)
            append_summary(compiled_text, heading=current_heading)
            section_text = []

        # for h, t in zip(headings, texts):
        #     print(h, t)
        # print('\n'.join(styles))

        if doc_file:
            if md_out:
                Path(doc_file).write_text('\n'.join(md_lines), encoding='utf-8')
            else:
                self._output_document.save(doc_file)
        if freeplane_file:
            self._output_midmapping.save(freeplane_file, encoding='utf-8')
        # if not doc_file and not freeplane_file:


def _odt_to_docx(odt_path: Path) -> Path:
    """Convert an ODT file to DOCX using LibreOffice headless. Returns the DOCX path."""
    tmpdir = Path(tempfile.mkdtemp())
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx",
         "--outdir", str(tmpdir), str(odt_path)],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed:\n{result.stderr}")
    converted = tmpdir / (odt_path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"Expected output not found: {converted}")
    return converted


class _MergedDocument:
    """Wraps multiple Document objects as a single iterable of paragraphs."""
    def __init__(self, docs, name="merged"):
        self.paragraphs = list(chain.from_iterable(d.paragraphs for d in docs))
        self.name = name


def process_files(input_paths: list,
        actions: dict,
        doc_path: Path = None, freeplane_path: Path = None, file_tag='outline'):
    """Process multiple files as a single merged document."""
    assert input_paths, 'no input files provided'

    docx_paths = []
    for p in input_paths:
        p = Path(p)
        assert p.exists(), f'input file not found: {p}'
        if p.suffix.lower() == '.odt':
            p = _odt_to_docx(p)
        docx_paths.append(p)

    docs = [Document(p) for p in docx_paths]
    merged_name = " + ".join(Path(p).stem for p in input_paths)

    merged = _MergedDocument(docs, name=merged_name)

    if doc_path and Path(doc_path).is_dir():
        doc_path = Path(doc_path, f'{merged_name}_{file_tag}.docx')
    if freeplane_path and Path(freeplane_path).is_dir():
        freeplane_path = Path(freeplane_path, f'{merged_name}_{file_tag}.mm')

    print(merged_name, ' --> ', doc_path, freeplane_path)

    outliner = Outliner.__new__(Outliner)
    outliner.file = docx_paths[0]
    outliner.name = merged_name
    outliner.document = merged
    outliner._output_document = Document()
    outliner._output_midmapping = None

    outliner.extract_marks(doc_file=doc_path, freeplane_file=freeplane_path, **actions)
    return doc_path, freeplane_path


def process_file(input_file_path: Path,
        actions: dict,
        doc_path: Path = None, freeplane_path: Path = None, file_tag='outline'):

    assert input_file_path.exists(), 'input file not found'

    if input_file_path.suffix.lower() == '.odt':
        input_file_path = _odt_to_docx(input_file_path)

    outliner = Outliner(input_file_path)

    if doc_path and doc_path.is_dir():
        doc_path = Path(doc_path, f'{input_file_path.stem}_{file_tag}.docx')

    if freeplane_path and freeplane_path.is_dir():
        freeplane_path = Path(freeplane_path,  f'{input_file_path.stem}_{file_tag}.mm')

    print(input_file_path, ' --> ', doc_path, freeplane_path)

    outliner.extract_marks(doc_file=doc_path, freeplane_file=freeplane_path, **actions)
    return doc_path, freeplane_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="""Reads a `docx` word document and outputs a new document that preserves document structure (headings) but replaces content with: 
+ An automatic summary (first detects language, then summarizes)
A compilation of bold/highlighted segments (they were annotated by the user first maybe because they represent important terms or ideas
    """)
    # positional
    parser.add_argument('input', metavar='input file', type=str,
                        help='input path (file or folder) for highlighting compilation or automatic summarization')
    parser.add_argument('-r', '--recursive', action='store_true',
                        help='Use it along a folder input and output: it makes available all files within subfolders')

    # conversion options
    parser.add_argument('-b', '--bold', action='store_true', help='filter bold content')
    parser.add_argument('-i', '--italic', action='store_true', help='filter italic content')
    parser.add_argument('-u', '--underline', action='store_true', help='filter underlined content')
    
    # OUTPUT
    # parser.add_argument('output', metavar='output file', type=str, help='output path (file or folder)')
    parser.add_argument('-d', '--docx', type=str, help='Specify word document output path')
    parser.add_argument('-f', '--freeplane', type=str, help='Specify freeplane document output path')

    parser.add_argument('-s', '--summary', action='store_true', help='summarize each section using an Ollama LLM')
    parser.add_argument('-m', '--model', type=str, default='llama3.2', help='Ollama model name (default: llama3.2)')
    # optional & mutually exclusive
    # parser.add_argument("-a", "--action", type=str, default='marks', choices=["marks", "summary"], help="Choose action to process document content")

    args = parser.parse_args()
    # print(args)

    assert args.docx or args.freeplane, 'A docx or freeplane output file/folder must be provided'
    input_path = Path(args.input)
    assert input_path.exists(), 'input file/folder not found'

    if input_path.is_file():
        files = [input_path]
    elif input_path.is_dir():
        # output must be a folder too
        if args.docx:
            assert Path(args.docx).is_dir()
        if args.freeplane:
            assert Path(args.freeplane).is_dir()

        # get input files
        glob = Path.rglob if args.recursive else Path.glob
        files = list(glob(input_path, '*.docx')) + list(glob(input_path, '*.odt'))
        assert files, 'No compatible files found in ' + str(input_path.resolve())

    if args.summary:
        from llm import run_task, DEFAULT_PROMPT
        model = args.model or 'llama3.2'
        actions = {'summarizer': lambda text: run_task(text, DEFAULT_PROMPT, model)}
        file_tag = 'summary'
    elif args.bold or args.italic or args.underline:
        actions = {'select_bold': args.bold, 'select_italic': args.italic, 'select_underline': args.underline}
        file_tag = 'annotation'
    else:
        raise SyntaxError('some arguments must be provided to indicate action on content: (-b -u | -s <lang>)')


    for file_path in files:
        process_file(file_path, actions,
                     Path(args.docx) if args.docx else None, Path(args.freeplane) if args.freeplane else None,
                     file_tag)
